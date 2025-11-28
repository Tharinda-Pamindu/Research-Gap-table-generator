import google.generativeai as genai
import PyPDF2
import docx
import pandas as pd
import io
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def extract_text_from_files(uploaded_files):
    """
    Extracts text from a list of uploaded files (PDF or DOCX).
    Returns a combined string of text and a list of filenames.
    """
    combined_text = ""
    filenames = []
    
    for uploaded_file in uploaded_files:
        filenames.append(uploaded_file.name)
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        try:
            if file_extension == 'pdf':
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page in pdf_reader.pages:
                    combined_text += page.extract_text() + "\n"
            elif file_extension in ['docx', 'doc']:
                doc = docx.Document(uploaded_file)
                for para in doc.paragraphs:
                    combined_text += para.text + "\n"
        except Exception as e:
            print(f"Error reading {uploaded_file.name}: {e}")
            
    return combined_text, filenames

def get_gemini_response(prompt, text_input, api_key):
    """
    Sends a prompt and text input to Google Gemini and returns the response.
    """
    if not api_key:
        return "Please provide a valid Google Gemini API Key."
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.5-flash')
        response = model.generate_content([prompt, text_input])
        return response.text
    except Exception as e:
        return f"Error accessing Gemini API: {str(e)}"

def generate_research_gap_table(text, api_key):
    """
    Generates a research gap table from the provided text using Gemini.
    Returns a Pandas DataFrame.
    """
    prompt = """
    You are an expert academic researcher. Analyze the provided research paper text and identify the research gaps.
    Create a comprehensive table summarizing the findings.
    
    The output MUST be a valid Markdown table with the following columns:
    | Reference | Year | Study Aim / Topic | Method / Approach | Data / Tools | Key Findings | Relevance to Project | Gaps / Notes | Research Gap / Limitations |
    
    IMPORTANT:
    - Do NOT include any introductory or concluding text.
    - Output ONLY the markdown table.
    - Ensure the table is well-formatted.
    - **Use IEEE Citation Style for the Reference column**: Format as "[1] Author(s), 'Title,' Journal/Conference, Year." Extract author names and titles from the text.
    - If author information is not available, use a descriptive reference like "[1] Study on [topic]".
    """
    
    response_text = get_gemini_response(prompt, text, api_key)
    
    # Parse markdown table to DataFrame
    try:
        lines = response_text.strip().split('\n')
        table_lines = [line.strip() for line in lines if '|' in line]
        
        if len(table_lines) < 2:
            return pd.DataFrame({"Error": ["Could not find a valid table in the response."], "Raw Response": [response_text]})

        # Extract headers from the first valid table line
        headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
        
        data = []
        # Iterate through the rest of the lines
        for line in table_lines[1:]:
            # Skip separator lines (e.g., |---|---|)
            if '---' in line:
                continue
                
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            
            # Handle row length mismatches
            if len(row) > 0:
                if len(row) == len(headers):
                    data.append(row)
                elif len(row) > len(headers):
                     data.append(row[:len(headers)])
                else:
                    # Pad with empty strings if short
                    row += [''] * (len(headers) - len(row))
                    data.append(row)
        
        if not data:
             return pd.DataFrame({"Error": ["Table found but no data parsed."], "Raw Response": [response_text]})
             
        df = pd.DataFrame(data, columns=headers)
        return df
    except Exception as e:
        return pd.DataFrame({"Error": [f"Failed to parse table: {str(e)}"], "Raw Response": [response_text]})

def answer_question(context, question, api_key):
    """
    Answers a user question based on the provided context.
    """
    prompt = f"""
    You are a helpful research assistant. Use the following context from research papers to answer the user's question.
    
    Context:
    {context[:30000]} # Limit context to avoid token limits if necessary, though Gemini handles large context well.
    
    Question: {question}
    
    Answer:
    """
    return get_gemini_response(prompt, "", api_key)

from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def create_pdf_download(df):
    """
    Creates a PDF file from the DataFrame with text wrapping and landscape orientation.
    """
    buffer = io.BytesIO()
    # Use A3 Landscape for more space (approx 1190 x 842 points)
    # Margins: 30 points each side
    from reportlab.lib.pagesizes import A3
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A3), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    normal_style = styles['Normal']
    normal_style.fontSize = 8 # Smaller font for table content
    
    elements.append(Paragraph("Research Gap Analysis", title_style))
    elements.append(Spacer(1, 12))
    
    # Convert DataFrame to list of lists, wrapping content in Paragraphs
    data = []
    
    # Headers
    headers = [Paragraph(f"<b>{col}</b>", normal_style) for col in df.columns]
    data.append(headers)
    
    # Rows
    for index, row in df.iterrows():
        row_data = []
        for item in row:
            # Clean text for ReportLab Paragraph
            text = str(item)
            
            # Truncate text if it's too long to prevent LayoutError
            # A3 height is ~840 pts. If a cell is > 800 pts, it crashes.
            # 1000 chars is a safe upper limit for 8pt font in a narrow column.
            if len(text) > 1000:
                text = text[:1000] + "... (truncated)"
            
            # Replace <br> with <br/> for ReportLab
            text = text.replace('<br>', '<br/>')
            
            # Wrap text in Paragraph to allow wrapping in table cells
            try:
                row_data.append(Paragraph(text, normal_style))
            except:
                # Fallback if parsing fails: strip tags and just show text
                import re
                clean_text = re.sub('<[^<]+?>', '', text)
                row_data.append(Paragraph(clean_text, normal_style))
                
        data.append(row_data)
    
    # Calculate column widths
    # Calculate usable width
    # A3 Landscape width is 420mm ~ 1190 points.
    # Margins are 30+30 = 60 points.
    usable_width = 1190 - 60 
    
    num_cols = len(df.columns)
    if num_cols > 0:
        # Distribute width evenly
        col_width = usable_width / num_cols
    else:
        col_width = 100
    
    # Create Table
    # repeatRows=1 repeats the header on new pages
    t = Table(data, colWidths=[col_width] * num_cols, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'), # Left align text
        ('VALIGN', (0, 0), (-1, -1), 'TOP'), # Top align text
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    
    elements.append(t)
    doc.build(elements)
    buffer.seek(0)
    return buffer

def create_docx_download(df):
    """
    Creates a DOCX file from the DataFrame.
    """
    doc = docx.Document()
    doc.add_heading('Research Gap Analysis', 0)
    
    # Add table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
    
    # Add data
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_literature_review(text, api_key):
    """
    Generates a literature review from the provided text using Gemini.
    """
    prompt = """
    You are an expert academic researcher. Write a comprehensive literature review based on the provided text from multiple research papers.
    
    Structure the review as follows:
    1.  **Introduction**: Briefly introduce the key themes and topics covered in the papers.
    2.  **Thematic Analysis**: Group the findings by common themes, methodologies, or debates. Compare and contrast the different studies.
    3.  **Critical Evaluation**: Discuss the strengths and weaknesses of the approaches used.
    4.  **Conclusion**: Summarize the state of the field and highlight any consensus or remaining gaps.
    5.  **References**: List the references in IEEE format corresponding to the citations used in the text.
    
    IMPORTANT:
    - Write in a formal, academic tone.
    - **Use IEEE Citation Style**: Use numeric citations in square brackets (e.g., [1], [2]) within the text.
    - Extract author names and titles from the text to create the References list.
    - Ensure the review is coherent and flows logically.
    """
    
    return get_gemini_response(prompt, text, api_key)

def create_review_docx(review_text):
    """
    Creates a DOCX file for the literature review.
    """
    doc = docx.Document()
    doc.add_heading('Literature Review', 0)
    
    # Split text by double newlines to create paragraphs
    paragraphs = review_text.split('\n\n')
    for para in paragraphs:
        doc.add_paragraph(para.strip())
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
